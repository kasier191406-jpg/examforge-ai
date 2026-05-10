from io import BytesIO
from tempfile import NamedTemporaryFile

import pandas as pd
from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Response, UploadFile, status
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.core.security import get_current_user, require_roles
from app.models import (
    AuditLog,
    ExamTemplate,
    Module,
    Question,
    QuestionPaper,
    QuestionPaperItem,
    Settings,
    Subject,
    User,
)
from app.schemas.catalog import (
    GeneratePaperRequest,
    ImportResult,
    ModuleCreate,
    ModuleRead,
    PaperQuestionRead,
    PaperRead,
    QuestionCreate,
    QuestionRead,
    QuestionUpdate,
    SettingsUpdate,
    SubjectCreate,
    SubjectRead,
    TemplateCreate,
)

router = APIRouter(tags=["catalog"])


def audit(db: Session, actor: str, action: str, entity: str, detail: str) -> None:
    db.add(AuditLog(actor=actor, action=action, entity=entity, detail=detail))


def question_to_read(question: Question) -> QuestionRead:
    return QuestionRead(
        id=question.id,
        subject_id=question.subject_id,
        module_id=question.module_id,
        subject_name=question.module.subject.name if question.module and question.module.subject else "",
        module_title=question.module.title if question.module else "",
        text=question.text,
        marks=question.marks,
        difficulty=question.difficulty,
        bloom_level=question.bloom_level,
        question_type=question.question_type,
        keywords=question.keywords,
        model_answer=question.model_answer,
        usage_count=question.usage_count,
        is_active=question.is_active,
    )


def paper_to_read(db: Session, paper: QuestionPaper) -> PaperRead:
    subject = db.get(Subject, paper.subject_id)
    item_rows = db.execute(
        select(QuestionPaperItem, Question)
        .join(Question, Question.id == QuestionPaperItem.question_id)
        .where(QuestionPaperItem.paper_id == paper.id)
        .order_by(QuestionPaperItem.sequence)
    ).all()
    return PaperRead(
        id=paper.id,
        title=paper.title,
        subject_id=paper.subject_id,
        subject_name=subject.name if subject else "",
        total_marks=paper.total_marks,
        duration_minutes=paper.duration_minutes,
        status=paper.status,
        created_at=paper.created_at.isoformat(),
        items=[
            PaperQuestionRead(
                id=item.id,
                question_id=question.id,
                section=item.section,
                sequence=item.sequence,
                marks=item.marks,
                text=question.text,
                difficulty=question.difficulty,
                bloom_level=question.bloom_level,
            )
            for item, question in item_rows
        ],
    )


@router.get("/users")
def list_users(
    db: Session = Depends(get_db),
    _: User = Depends(require_roles("admin")),
) -> list[dict[str, object]]:
    users = db.scalars(select(User).order_by(User.username)).all()
    return [
        {
            "id": user.id,
            "username": user.username,
            "email": user.email,
            "full_name": user.full_name,
            "role": user.role,
            "is_active": user.is_active,
        }
        for user in users
    ]


@router.get("/subjects", response_model=list[SubjectRead])
def list_subjects(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> list[SubjectRead]:
    subjects = db.scalars(select(Subject).order_by(Subject.name)).all()
    return [
        SubjectRead(
            id=subject.id,
            code=subject.code,
            name=subject.name,
            description=subject.description,
            modules=len(subject.modules),
        )
        for subject in subjects
    ]


@router.post("/subjects", response_model=SubjectRead, status_code=status.HTTP_201_CREATED)
def create_subject(
    payload: SubjectCreate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> SubjectRead:
    subject = Subject(**payload.model_dump())
    db.add(subject)
    audit(db, user.username, "create", "subject", payload.name)
    db.commit()
    db.refresh(subject)
    return SubjectRead(id=subject.id, modules=0, **payload.model_dump())


@router.get("/modules", response_model=list[ModuleRead])
def list_modules(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> list[ModuleRead]:
    modules = db.scalars(select(Module).order_by(Module.subject_id, Module.order_index)).all()
    return [ModuleRead(id=module.id, subject_id=module.subject_id, title=module.title, order_index=module.order_index) for module in modules]


@router.post("/modules", response_model=ModuleRead, status_code=status.HTTP_201_CREATED)
def create_module(
    payload: ModuleCreate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> ModuleRead:
    if db.get(Subject, payload.subject_id) is None:
        raise HTTPException(status_code=404, detail="Subject not found")
    module = Module(**payload.model_dump())
    db.add(module)
    audit(db, user.username, "create", "module", payload.title)
    db.commit()
    db.refresh(module)
    return ModuleRead(id=module.id, subject_id=module.subject_id, title=module.title, order_index=module.order_index)


@router.get("/questions", response_model=list[QuestionRead])
def list_questions(
    limit: int = 200,
    search: str = "",
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> list[QuestionRead]:
    statement = select(Question).order_by(Question.id.desc()).limit(limit)
    if search:
        statement = select(Question).where(Question.text.ilike(f"%{search}%")).order_by(Question.id.desc()).limit(limit)
    return [question_to_read(question) for question in db.scalars(statement).all()]


@router.post("/questions", response_model=QuestionRead, status_code=status.HTTP_201_CREATED)
def create_question(
    payload: QuestionCreate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> QuestionRead:
    if db.get(Subject, payload.subject_id) is None or db.get(Module, payload.module_id) is None:
        raise HTTPException(status_code=404, detail="Subject or module not found")
    question = Question(**payload.model_dump(), usage_count=0)
    db.add(question)
    audit(db, user.username, "create", "question", question.text[:120])
    db.commit()
    db.refresh(question)
    return question_to_read(question)


@router.put("/questions/{question_id}", response_model=QuestionRead)
def update_question(
    question_id: int,
    payload: QuestionUpdate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> QuestionRead:
    question = db.get(Question, question_id)
    if question is None:
        raise HTTPException(status_code=404, detail="Question not found")
    for field, value in payload.model_dump(exclude_unset=True).items():
        setattr(question, field, value)
    audit(db, user.username, "update", "question", str(question_id))
    db.commit()
    db.refresh(question)
    return question_to_read(question)


@router.delete("/questions/{question_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_question(
    question_id: int,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> Response:
    question = db.get(Question, question_id)
    if question is None:
        raise HTTPException(status_code=404, detail="Question not found")
    db.delete(question)
    audit(db, user.username, "delete", "question", str(question_id))
    db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@router.get("/templates")
def list_templates(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> list[dict[str, object]]:
    templates = db.scalars(select(ExamTemplate).order_by(ExamTemplate.id)).all()
    return [
        {
            "id": template.id,
            "name": template.name,
            "duration_minutes": template.duration_minutes,
            "total_marks": template.total_marks,
            "sections": template.sections,
            "instructions": template.instructions,
        }
        for template in templates
    ]


@router.post("/templates", status_code=status.HTTP_201_CREATED)
def create_template(
    payload: TemplateCreate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> dict[str, object]:
    template = ExamTemplate(**payload.model_dump())
    db.add(template)
    audit(db, user.username, "create", "template", template.name)
    db.commit()
    db.refresh(template)
    return {"id": template.id, **payload.model_dump()}


@router.get("/papers", response_model=list[PaperRead])
def list_papers(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> list[PaperRead]:
    papers = db.scalars(select(QuestionPaper).order_by(QuestionPaper.created_at.desc())).all()
    return [paper_to_read(db, paper) for paper in papers]


@router.get("/papers/{paper_id}", response_model=PaperRead)
def get_paper(paper_id: int, db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> PaperRead:
    paper = db.get(QuestionPaper, paper_id)
    if paper is None:
        raise HTTPException(status_code=404, detail="Paper not found")
    return paper_to_read(db, paper)


def choose_questions(candidates: list[Question], total_marks: int) -> list[Question]:
    dp: dict[int, list[Question]] = {0: []}
    for question in candidates:
        for marks, selected in list(dp.items()):
            next_marks = marks + question.marks
            if next_marks <= total_marks and next_marks not in dp:
                dp[next_marks] = [*selected, question]
        if total_marks in dp:
            return dp[total_marks]
    return dp.get(total_marks, [])


@router.post("/generator", response_model=PaperRead, status_code=status.HTTP_201_CREATED)
def generate_paper(
    payload: GeneratePaperRequest,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> PaperRead:
    statement = select(Question).where(Question.subject_id == payload.subject_id, Question.is_active.is_(True))
    if payload.module_ids:
        statement = statement.where(Question.module_id.in_(payload.module_ids))
    candidates = db.scalars(statement.order_by(Question.usage_count, Question.id)).all()
    selected = choose_questions(candidates, payload.total_marks)
    if sum(question.marks for question in selected) != payload.total_marks:
        raise HTTPException(status_code=400, detail="Unable to satisfy exact total marks with available questions")

    paper = QuestionPaper(
        title=payload.title,
        subject_id=payload.subject_id,
        total_marks=payload.total_marks,
        duration_minutes=payload.duration_minutes,
        status="published",
        generated_by=user.id,
    )
    db.add(paper)
    db.flush()
    for index, question in enumerate(selected, start=1):
        question.usage_count += 1
        db.add(
            QuestionPaperItem(
                paper_id=paper.id,
                question_id=question.id,
                section=chr(65 + ((index - 1) % payload.sections)),
                sequence=index,
                marks=question.marks,
            )
        )
    audit(db, user.username, "generate", "paper", f"{paper.title} ({payload.total_marks} marks)")
    db.commit()
    db.refresh(paper)
    return paper_to_read(db, paper)


@router.get("/papers/{paper_id}/export/pdf")
def export_pdf(paper_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)) -> Response:
    paper = get_paper(paper_id, db, user)
    settings = db.scalar(select(Settings).order_by(Settings.id).limit(1))
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - inch
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawCentredString(width / 2, y, settings.institution_name if settings else "ExamForge AI")
    y -= 24
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawCentredString(width / 2, y, paper.title)
    y -= 20
    pdf.setFont("Helvetica", 10)
    pdf.drawCentredString(width / 2, y, f"{paper.subject_name} | {paper.duration_minutes} minutes | {paper.total_marks} marks")
    y -= 36
    for item in paper.items:
        if y < inch:
            pdf.showPage()
            y = height - inch
        pdf.setFont("Helvetica-Bold", 10)
        pdf.drawString(inch, y, f"Q{item.sequence}. [{item.marks}]")
        pdf.setFont("Helvetica", 10)
        text = pdf.beginText(inch + 52, y)
        text.textLines(item.text[:420])
        pdf.drawText(text)
        y -= 44
    pdf.save()
    audit(db, user.username, "export", "paper_pdf", str(paper_id))
    db.commit()
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="paper-{paper_id}.pdf"'},
    )


@router.get("/papers/{paper_id}/export/docx")
def export_docx(paper_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)) -> Response:
    paper = get_paper(paper_id, db, user)
    document = Document()
    document.add_heading("Tech University", level=1)
    document.add_heading(paper.title, level=2)
    document.add_paragraph(f"{paper.subject_name} | {paper.duration_minutes} minutes | {paper.total_marks} marks")
    for item in paper.items:
        document.add_paragraph(f"Q{item.sequence}. {item.text} [{item.marks}]", style=None)
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        document.save(tmp.name)
        tmp.seek(0)
        content = tmp.read()
    audit(db, user.username, "export", "paper_docx", str(paper_id))
    db.commit()
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="paper-{paper_id}.docx"'},
    )


@router.post("/import/questions", response_model=ImportResult)
async def import_questions(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin", "faculty")),
) -> ImportResult:
    content = await file.read()
    try:
        dataframe = pd.read_excel(BytesIO(content)) if file.filename.endswith((".xlsx", ".xls")) else pd.read_csv(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Unable to read import file: {exc}") from exc

    required = {"subject", "module", "text", "marks", "difficulty", "bloom_level"}
    missing = required - set(dataframe.columns)
    if missing:
        raise HTTPException(status_code=400, detail=f"Missing columns: {', '.join(sorted(missing))}")

    imported = 0
    skipped = 0
    errors: list[str] = []
    for row_number, row in dataframe.fillna("").iterrows():
        try:
            subject_name = str(row["subject"]).strip()
            module_title = str(row["module"]).strip()
            text = str(row["text"]).strip()
            if not subject_name or not module_title or len(text) < 10:
                skipped += 1
                errors.append(f"Row {row_number + 2}: subject, module, and text are required")
                continue
            duplicate = db.scalar(select(Question).where(Question.text == text))
            if duplicate:
                skipped += 1
                continue
            subject = db.scalar(select(Subject).where(Subject.name == subject_name))
            if subject is None:
                code = subject_name.upper().replace(" ", "-")[:30]
                subject = Subject(code=code, name=subject_name, description="Imported subject")
                db.add(subject)
                db.flush()
            module = db.scalar(select(Module).where(Module.subject_id == subject.id, Module.title == module_title))
            if module is None:
                module = Module(subject_id=subject.id, title=module_title, order_index=1)
                db.add(module)
                db.flush()
            db.add(
                Question(
                    subject_id=subject.id,
                    module_id=module.id,
                    text=text,
                    marks=int(row["marks"]),
                    difficulty=str(row["difficulty"]).strip(),
                    bloom_level=str(row["bloom_level"]).strip(),
                    question_type=str(row.get("question_type", "descriptive") or "descriptive"),
                    keywords=str(row.get("keywords", "")),
                    model_answer=str(row.get("model_answer", "")),
                    usage_count=0,
                    is_active=True,
                )
            )
            imported += 1
        except Exception as exc:
            skipped += 1
            errors.append(f"Row {row_number + 2}: {exc}")
    audit(db, user.username, "import", "questions", f"imported={imported}, skipped={skipped}")
    db.commit()
    return ImportResult(imported=imported, skipped=skipped, errors=errors)


@router.get("/analytics")
def analytics(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> dict[str, object]:
    bloom_rows = db.execute(select(Question.bloom_level, func.count()).group_by(Question.bloom_level)).all()
    difficulty_rows = db.execute(select(Question.difficulty, func.count()).group_by(Question.difficulty)).all()
    module_rows = db.execute(select(Module.title, func.count(Question.id)).join(Question).group_by(Module.title).limit(10)).all()
    used = db.scalars(select(Question).order_by(Question.usage_count.desc()).limit(8)).all()
    papers = db.scalars(select(QuestionPaper).order_by(QuestionPaper.created_at)).all()
    return {
        "bloom_distribution": [{"name": name, "value": count} for name, count in bloom_rows],
        "difficulty_breakdown": [{"name": name, "value": count} for name, count in difficulty_rows],
        "module_coverage": [{"name": name, "value": count} for name, count in module_rows],
        "generation_trends": [{"name": paper.created_at.strftime("%m-%d"), "value": paper.total_marks} for paper in papers],
        "most_used_questions": [{"id": question.id, "text": question.text, "usage_count": question.usage_count} for question in used],
    }


@router.get("/settings")
def get_settings(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> dict[str, object]:
    row = db.scalar(select(Settings).order_by(Settings.id).limit(1))
    if row is None:
        row = Settings(institution_name="Tech University", similarity_threshold=80, default_instructions="")
        db.add(row)
        db.commit()
        db.refresh(row)
    return {
        "institution_name": row.institution_name,
        "logo_url": row.logo_url,
        "similarity_threshold": row.similarity_threshold / 100,
        "default_instructions": row.default_instructions,
    }


@router.put("/settings")
def update_settings(
    payload: SettingsUpdate,
    db: Session = Depends(get_db),
    user: User = Depends(require_roles("admin")),
) -> dict[str, object]:
    row = db.scalar(select(Settings).order_by(Settings.id).limit(1))
    if row is None:
        row = Settings()
        db.add(row)
    row.institution_name = payload.institution_name
    row.logo_url = payload.logo_url
    row.similarity_threshold = int(payload.similarity_threshold * 100)
    row.default_instructions = payload.default_instructions
    audit(db, user.username, "update", "settings", row.institution_name)
    db.commit()
    return get_settings(db, user)


@router.get("/audit-logs")
def list_audit_logs(db: Session = Depends(get_db), _: User = Depends(get_current_user)) -> list[dict[str, object]]:
    rows = db.scalars(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(50)).all()
    return [
        {
            "id": row.id,
            "actor": row.actor,
            "action": row.action,
            "entity": row.entity,
            "detail": row.detail,
            "created_at": row.created_at.isoformat(),
        }
        for row in rows
    ]
